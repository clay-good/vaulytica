"""
Document Ingestion Module for Security Questionnaire Agent

Extracts text and data from various file formats:
- PDF (PyPDF2, pdfplumber)
- DOCX (python-docx)
- Excel (openpyxl, pandas)
- CSV (pandas)
- Markdown (plain text)
- TXT (plain text)

Version: 1.0.0
Author: Vaulytica Team
"""

import asyncio
import hashlib
import mimetypes
import re
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import logging

logger = logging.getLogger(__name__)


class DocumentFormat(str, Enum):
    """Supported document formats"""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    CSV = "csv"
    MARKDOWN = "md"
    TEXT = "txt"
    UNKNOWN = "unknown"


class DocumentType(str, Enum):
    """Types of security documents"""
    SECURITY_POLICY = "security_policy"
    SOP = "sop"
    ARCHITECTURE = "architecture"
    COMPLIANCE_DOC = "compliance_doc"
    QUESTIONNAIRE = "questionnaire"
    RUNBOOK = "runbook"
    VENDOR_CONTRACT = "vendor_contract"
    EMPLOYEE_HANDBOOK = "employee_handbook"
    AUDIT_REPORT = "audit_report"
    OTHER = "other"


@dataclass
class ExtractedDocument:
    """Document extracted from a file"""
    document_id: str
    file_path: str
    file_name: str
    file_format: DocumentFormat
    document_type: DocumentType
    title: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    tags: List[str] = field(default_factory=list)
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    extracted_at: datetime = field(default_factory=datetime.utcnow)
    file_size_bytes: Optional[int] = None
    checksum: Optional[str] = None


class DocumentIngestionModule:
    """
    Document ingestion module for extracting text from various file formats.

    Supports:
    - PDF: PyPDF2 (fallback to pdfplumber for complex PDFs)
    - DOCX: python-docx
    - Excel: openpyxl + pandas
    - CSV: pandas
    - Markdown: plain text
    - TXT: plain text
    """

    def __init__(self):
        self.supported_formats = {
            ".pdf": DocumentFormat.PDF,
            ".docx": DocumentFormat.DOCX,
            ".xlsx": DocumentFormat.XLSX,
            ".xls": DocumentFormat.XLSX,
            ".csv": DocumentFormat.CSV,
            ".md": DocumentFormat.MARKDOWN,
            ".txt": DocumentFormat.TEXT
        }
        logger.info("DocumentIngestionModule initialized")

    async def ingest_file(
        self,
        file_path: Union[str, Path],
        document_type: DocumentType = DocumentType.OTHER,
        title: Optional[str] = None,
        tags: Optional[List[str]] = None
    ) -> ExtractedDocument:
        """
        Ingest a file and extract its content.

        Args:
            file_path: Path to the file
            document_type: Type of document
            title: Optional title (defaults to filename)
            tags: Optional tags for categorization

        Returns:
            ExtractedDocument with extracted content

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect format
        file_format = self._detect_format(file_path)
        if file_format == DocumentFormat.UNKNOWN:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")

        logger.info(f"Ingesting {file_format.value} file: {file_path.name}")

        # Extract content based on format
        if file_format == DocumentFormat.PDF:
            content, metadata = await self._extract_pdf(file_path)
        elif file_format == DocumentFormat.DOCX:
            content, metadata = await self._extract_docx(file_path)
        elif file_format == DocumentFormat.XLSX:
            content, metadata = await self._extract_excel(file_path)
        elif file_format == DocumentFormat.CSV:
            content, metadata = await self._extract_csv(file_path)
        elif file_format == DocumentFormat.MARKDOWN:
            content, metadata = await self._extract_markdown(file_path)
        elif file_format == DocumentFormat.TEXT:
            content, metadata = await self._extract_text(file_path)
        else:
            raise ValueError(f"Unsupported format: {file_format}")

        # Clean content
        content = self._clean_text(content)

        # Generate document ID
        document_id = self._generate_document_id(file_path, content)

        # Calculate file stats
        file_size = file_path.stat().st_size
        word_count = len(content.split())
        checksum = self._calculate_checksum(file_path)

        # Create extracted document
        extracted_doc = ExtractedDocument(
            document_id=document_id,
            file_path=str(file_path),
            file_name=file_path.name,
            file_format=file_format,
            document_type=document_type,
            title=title or file_path.stem,
            content=content,
            metadata=metadata,
            tags=tags or [],
            page_count=metadata.get("page_count"),
            word_count=word_count,
            file_size_bytes=file_size,
            checksum=checksum
        )

        logger.info(f"Extracted {word_count} words from {file_path.name}")
        return extracted_doc

    def _detect_format(self, file_path: Path) -> DocumentFormat:
        """Detect file format from extension"""
        suffix = file_path.suffix.lower()
        return self.supported_formats.get(suffix, DocumentFormat.UNKNOWN)

    async def _extract_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from PDF using PyPDF2 (fallback to pdfplumber)"""
        try:
            import PyPDF2

            content_parts = []
            metadata = {}

            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                metadata["page_count"] = len(reader.pages)

                # Extract metadata
                if reader.metadata:
                    metadata["pdf_title"] = reader.metadata.get("/Title", "")
                    metadata["pdf_author"] = reader.metadata.get("/Author", "")
                    metadata["pdf_subject"] = reader.metadata.get("/Subject", "")

                # Extract text from each page
                for page_num, page in enumerate(reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text:
                            content_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num}: {e}")

            content = "\n\n".join(content_parts)

            # If PyPDF2 extraction is poor, try pdfplumber
            if len(content.strip()) < 100:
                logger.info("PyPDF2 extraction poor, trying pdfplumber...")
                content, metadata = await self._extract_pdf_pdfplumber(file_path)

            return content, metadata

        except ImportError:
            logger.warning("PyPDF2 not installed, trying pdfplumber...")
            return await self._extract_pdf_pdfplumber(file_path)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return "", {"error": str(e)}

    async def _extract_pdf_pdfplumber(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from PDF using pdfplumber (better for complex PDFs)"""
        try:
            import pdfplumber

            content_parts = []
            metadata = {}

            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)

                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)

            content = "\n\n".join(content_parts)
            return content, metadata

        except ImportError:
            logger.error("pdfplumber not installed. Install with: pip install pdfplumber")
            return "", {"error": "pdfplumber not installed"}
        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {e}")
            return "", {"error": str(e)}

    async def _extract_docx(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from DOCX using python-docx"""
        try:
            from docx import Document

            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    tables_text.append(row_text)

            content = "\n\n".join(paragraphs)
            if tables_text:
                content += "\n\n" + "\n".join(tables_text)

            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            }

            return content, metadata

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return "", {"error": "python-docx not installed"}
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return "", {"error": str(e)}

    async def _extract_excel(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from Excel using pandas"""
        try:
            import pandas as pd

            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names

            content_parts = []
            metadata = {"sheet_count": len(sheet_names), "sheets": []}

            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # Convert DataFrame to text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False)
                content_parts.append(sheet_text)

                metadata["sheets"].append({
                    "name": sheet_name,
                    "rows": len(df),
                    "columns": len(df.columns)
                })

            content = "\n\n".join(content_parts)
            return content, metadata

        except ImportError:
            logger.error("pandas/openpyxl not installed. Install with: pip install pandas openpyxl")
            return "", {"error": "pandas/openpyxl not installed"}
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            return "", {"error": str(e)}

    async def _extract_csv(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from CSV using pandas"""
        try:

            df = pd.read_csv(file_path)
            content = df.to_string(index=False)

            metadata = {
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": list(df.columns)
            }

            return content, metadata

        except ImportError:
            logger.error("pandas not installed. Install with: pip install pandas")
            # Fallback to basic CSV reading
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return content, {}
        except Exception as e:
            logger.error(f"CSV extraction failed: {e}")
            return "", {"error": str(e)}

    async def _extract_markdown(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Count headers
            header_count = len(re.findall(r'^#+\s', content, re.MULTILINE))

            metadata = {"header_count": header_count}
            return content, metadata

        except Exception as e:
            logger.error(f"Markdown extraction failed: {e}")
            return "", {"error": str(e)}

    async def _extract_text(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return content, {}
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return "", {"error": str(e)}

    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)

        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)

        return text.strip()

    def _generate_document_id(self, file_path: Path, content: str) -> str:
        """Generate unique document ID"""
        unique_string = f"{file_path.name}_{content[:100]}_{datetime.utcnow().isoformat()}"
        return hashlib.sha256(unique_string.encode()).hexdigest()[:16]

    def _calculate_checksum(self, file_path: Path) -> str:
        """Calculate file checksum (SHA256)"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()


# Singleton instance
_document_ingestion_module: Optional[DocumentIngestionModule] = None


def get_document_ingestion_module() -> DocumentIngestionModule:
    """Get singleton instance of DocumentIngestionModule"""
    global _document_ingestion_module
    if _document_ingestion_module is None:
        _document_ingestion_module = DocumentIngestionModule()
    return _document_ingestion_module

